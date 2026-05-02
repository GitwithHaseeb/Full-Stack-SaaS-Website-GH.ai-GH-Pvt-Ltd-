"""
Convert docs/IEEE-GH-AI-Project-Report.md to a Word document (.docx).

Requires: pip install python-docx

Usage (from repository root):
    python docs/generate_ieee_report.py

Output:
    docs/IEEE-GH-AI-Project-Report.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def strip_inline_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def main() -> None:
    base = Path(__file__).resolve().parent
    md_path = base / "IEEE-GH-AI-Project-Report.md"
    out_path = base / "IEEE-GH-AI-Project-Report.docx"
    if not md_path.is_file():
        raise SystemExit(f"Missing markdown report: {md_path}")

    doc = Document()
    doc.core_properties.title = "GH.ai — Technical Report (IEEE-style)"
    doc.core_properties.author = "GH Pvt Ltd"

    in_fence = False
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            p = doc.add_paragraph(line)
            for r in p.runs:
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            continue

        if stripped == "---":
            doc.add_page_break()
            continue
        if not stripped:
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif stripped.startswith("|") and "|" in stripped[1:]:
            p = doc.add_paragraph(strip_inline_bold(line))
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10)
        else:
            doc.add_paragraph(strip_inline_bold(line))

    doc.save(out_path)
    print("Wrote", out_path.name)


if __name__ == "__main__":
    main()
